import os
import sys

try:
    import docx
except ImportError:
    print("python-docx not installed. Installing...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Set margins to 1 inch
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Document Title
title = doc.add_paragraph()
title_run = title.add_run("PETA JALAN (ROADMAP) PROYEK WEBSITE HIQMA")
title_run.font.name = 'Plus Jakarta Sans'
title_run.font.size = Pt(18)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(15, 45, 89) # Navy
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph()
subtitle_run = subtitle.add_run("Target Peluncuran: 14 Agustus 2026 | Sisa Waktu Kerja: ~8 Minggu")
subtitle_run.font.name = 'Plus Jakarta Sans'
subtitle_run.font.size = Pt(11)
subtitle_run.font.italic = True
subtitle_run.font.color.rgb = RGBColor(212, 175, 55) # Gold
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph() # Spacing

# Introduction
p = doc.add_paragraph()
p_run = p.add_run("Dengan sisa waktu sekitar 8 minggu dari sekarang (17 Juni 2026), jadwal pengerjaan Express, MySQL, dan HTML/CSS/JS ini sangat realistis untuk diselesaikan dengan kualitas tinggi tanpa terburu-buru. Berikut adalah rencana distribusi kerja mingguan yang terstruktur:")
p_run.font.name = 'Plus Jakarta Sans'
p_run.font.size = Pt(11)

# Weekly Timeline
weeks = [
    ("Minggu 1-2 (17 Juni – 30 Juni)", "Integrasi Fitur Booking (Fase 4)", 
     "Fokus Utama: Full-stack integration formulir pemesanan.\n\nDeliverables:\n- Endpoint API POST /api/bookings di backend.\n- Logika generate nomor booking otomatis (HQM-YYYYMMDD-XXXX).\n- Integrasi data booking dengan relasi many-to-many ke layanan talent.\n- Uji koneksi frontend ke API Express & penyimpanan data ke database MySQL di Docker."),
    ("Minggu 3-4 (1 Juli – 14 Juli)", "Portal Anggota & Autentikasi (Fase 5 - Bagian 1)", 
     "Fokus Utama: Autentikasi login & Dashboard Anggota.\n\nDeliverables:\n- Pembuatan UI halaman Login Anggota.\n- API Register & Login dengan enkripsi kata sandi menggunakan bcrypt dan token JSON Web Token (JWT).\n- Dashboard Anggota: Menampilkan profil biodata pribadi dan data divisi resmi yang diikuti (sesuai data SK Pelantikan)."),
    ("Minggu 5-6 (15 Juli – 28 Juli)", "Dashboard Admin & CMS (Fase 5 - Bagian 2)", 
     "Fokus Utama: Manajemen data oleh admin.\n\nDeliverables:\n- Dashboard Admin: Halaman ringkasan statistik (jumlah booking, jumlah anggota).\n- Panel Booking: Fitur mengubah status booking (Menunggu/Diproses/Diterima/Ditolak/Selesai) disertai pengisian catatan admin.\n- CMS Sederhana: Fitur CRUD (Create, Read, Update, Delete) berita untuk halaman Publikasi dan unggah tautan dokumentasi untuk Galeri."),
    ("Minggu 7 (29 Juli – 4 Agustus)", "Pengujian Sistem & Keamanan (Fase 6 - Bagian 1)", 
     "Fokus Utama: Testing fungsionalitas & proteksi sistem.\n\nDeliverables:\n- Uji responsivitas menyeluruh di berbagai ukuran perangkat.\n- Pengujian celah keamanan dasar (SQL Injection prevention menggunakan parameter query, XSS filter pada form input).\n- Perbaikan bug dan validasi input (contoh: memastikan format email dan no. handphone pemesan valid)."),
    ("Minggu 8 (5 Agustus – 14 Agustus)", "Deployment & Peluncuran (Fase 6 - Bagian 2)", 
     "Fokus Utama: Go-live ke VPS server produksi.\n\nDeliverables:\n- Konfigurasi berkas produksi .env yang aman.\n- Setup Docker database MySQL dan aplikasi Express di server VPS.\n- Penghubungan domain resmi hiqmauinjkt.com.\n- Instalasi SSL gratis (Let's Encrypt) demi protokol HTTPS yang aman.\n- 14 Agustus 2026: Website resmi diluncurkan secara resmi.")
]

for week_title, focus, details in weeks:
    h = doc.add_paragraph()
    h_run = h.add_run(f"\n{week_title}: {focus}")
    h_run.font.name = 'Plus Jakarta Sans'
    h_run.font.size = Pt(13)
    h_run.font.bold = True
    h_run.font.color.rgb = RGBColor(15, 45, 89)
    
    d = doc.add_paragraph()
    d_run = d.add_run(details)
    d_run.font.name = 'Plus Jakarta Sans'
    d_run.font.size = Pt(10.5)

# Mitigasi Risiko
h_risk = doc.add_paragraph()
h_risk_run = h_risk.add_run("\nRencana Mitigasi Risiko (Risk Management)")
h_risk_run.font.name = 'Plus Jakarta Sans'
h_risk_run.font.size = Pt(14)
h_risk_run.font.bold = True
h_risk_run.font.color.rgb = RGBColor(15, 45, 89)

r = doc.add_paragraph()
r_run = r.add_run(
    "1. Risiko: Docker MySQL Error di Komputer Lokal\n"
    "   Mitigasi: Jika Docker Desktop terus bermasalah dengan DNS koneksi, kita dapat menggunakan instalasi MySQL lokal biasa (seperti XAMPP/Laragon) untuk tahap development, dan baru bermigrasi ke Docker saat dideploy di VPS.\n\n"
    "2. Risiko: Keterlambatan Konten Publikasi\n"
    "   Mitigasi: Menggunakan data tiruan (dummy data) yang realistis terlebih dahulu agar sistem database berfungsi penuh, sembari menunggu data ril disiapkan oleh pengurus divisi Humas/Media."
)
r_run.font.name = 'Plus Jakarta Sans'
r_run.font.size = Pt(10.5)

output_path = "Project_Roadmap_HIQMA.docx"
doc.save(output_path)
print(f"Word document saved successfully to: {os.path.abspath(output_path)}")
